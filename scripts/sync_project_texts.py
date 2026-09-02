from __future__ import annotations

import csv
import hashlib
import io
import json
import re
import unicodedata
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path

import requests
from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data" / "legislative" / "2026"
OUT = DATA / "topics"
TEXTS_PATH = OUT / "project_texts.jsonl"
INDEX_PATH = OUT / "project_text_index.csv"
DIAG_PATH = OUT / "text_sync_diagnostics.json"
MAX_PDF_PAGES = 25
MAX_TEXT_CHARS = 120_000

HTTP = requests.Session()
HTTP.headers.update({
    "User-Agent": "Mozilla/5.0 conoce-a-tu-parlamentario/project-text-sync",
    "Accept-Language": "es-CL,es;q=0.9",
})


def read_csv(path: Path) -> list[dict]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def norm(value: str) -> str:
    text = unicodedata.normalize("NFD", str(value or "").lower())
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    return re.sub(r"\s+", " ", text).strip()


def document_priority(event: dict) -> int:
    label = norm(event.get("subetapa", ""))
    if not (event.get("documento_url") or "").strip():
        return 99
    if "texto del mensaje" in label:
        return 0
    if "texto del proyecto" in label:
        return 1
    if "ingreso de proyecto" in label:
        return 2
    return 50


def classify_document(event: dict) -> str:
    label = norm(event.get("subetapa", ""))
    if "texto del mensaje" in label:
        return "texto_mensaje"
    if "texto del proyecto" in label:
        return "texto_proyecto"
    if "ingreso de proyecto" in label:
        return "ingreso_proyecto"
    return "otro"


def extract_pdf(content: bytes) -> tuple[str, int, int]:
    reader = PdfReader(io.BytesIO(content))
    total = len(reader.pages)
    examined = min(total, MAX_PDF_PAGES)
    parts = []
    for page in reader.pages[:examined]:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return "\n".join(parts), total, examined


def extract_docx(content: bytes) -> tuple[str, int]:
    document = Document(io.BytesIO(content))
    pieces = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            pieces.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    return "\n".join(pieces), len(document.paragraphs)


def extract_html(content: bytes, encoding: str | None) -> str:
    text = content.decode(encoding or "utf-8", errors="replace")
    soup = BeautifulSoup(text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    return soup.get_text("\n", strip=True)


def clean_text(text: str) -> str:
    raw = re.sub(r"[\t\r]+", " ", text or "")
    raw = re.sub(r" +", " ", raw)
    raw = re.sub(r"\n{3,}", "\n\n", raw).strip()

    # Varias mociones de la Cámara comienzan con una ficha administrativa que
    # enumera todas las comisiones posibles. Solo recortamos el prefijo cuando
    # hay señales claras de esa plantilla y encontramos un marcador sustantivo.
    upper = raw.upper()
    boilerplate_clues = sum(
        clue in upper
        for clue in (
            "PERIODO LEGISLATIVO",
            "LEGISLATURA",
            "PRIMER TRÁMITE",
            "01.- AGRICULTURA",
            "DESTINACIÓN",
        )
    )
    markers = (
        "FUNDAMENTOS",
        "ANTECEDENTES",
        "EXPOSICIÓN DE MOTIVOS",
        "EXPOSICION DE MOTIVOS",
        "IDEA MATRIZ",
        "IDEA MATRIZ O FUNDAMENTAL",
    )
    if boilerplate_clues >= 2:
        positions = [upper.find(marker) for marker in markers if upper.find(marker) >= 0]
        if positions:
            cut = min(positions)
            if 0 < cut < min(len(raw), 12_000):
                raw = raw[cut:]

    # Si quedó una enumeración completa de comisiones antes del contenido,
    # la reducimos de forma conservadora hasta un marcador sustantivo.
    raw = re.sub(
        r"01\.-\s*AGRICULTURA.*?(?=(?:FUNDAMENTOS|ANTECEDENTES|EXPOSICI[ÓO]N DE MOTIVOS|IDEA MATRIZ))",
        "",
        raw,
        flags=re.IGNORECASE | re.DOTALL,
    )
    raw = re.sub(r"\s+", " ", raw).strip()
    return raw[:MAX_TEXT_CHARS]


def quality(chars: int) -> str:
    if chars >= 2000:
        return "rica"
    if chars >= 500:
        return "utilizable"
    if chars > 0:
        return "escasa"
    return "sin_texto"


def fetch_document(url: str) -> dict:
    response = HTTP.get(url, timeout=60, allow_redirects=True)
    response.raise_for_status()
    content = response.content
    ctype = (response.headers.get("Content-Type") or "").lower()
    method = ""
    total_pages = 0
    examined = 0
    raw_text = ""

    if content.lstrip().startswith(b"%PDF") or "application/pdf" in ctype:
        method = "pdf_text"
        raw_text, total_pages, examined = extract_pdf(content)
    elif "officedocument.wordprocessingml.document" in ctype:
        method = "docx_text"
        raw_text, examined = extract_docx(content)
    elif "html" in ctype or content.lstrip().startswith(b"<"):
        method = "html_text"
        raw_text = extract_html(content, response.encoding)
    else:
        method = "unsupported"

    cleaned = clean_text(raw_text)
    return {
        "resolved_url": response.url,
        "mime_type": ctype.split(";", 1)[0],
        "extraction_method": method,
        "pages_total": total_pages,
        "units_examined": examined,
        "raw_chars": len(re.sub(r"\s+", " ", raw_text or "").strip()),
        "cleaned_chars": len(cleaned),
        "text_quality": quality(len(cleaned)),
        "cleaned_text": cleaned,
        "text_sha256": hashlib.sha256(cleaned.encode("utf-8")).hexdigest() if cleaned else "",
    }


def load_existing() -> dict[str, dict]:
    if not TEXTS_PATH.exists():
        return {}
    result = {}
    with TEXTS_PATH.open("r", encoding="utf-8") as handle:
        for line in handle:
            if line.strip():
                row = json.loads(line)
                if row.get("boletin"):
                    result[row["boletin"]] = row
    return result


def main() -> None:
    projects = read_csv(DATA / "projects.csv")
    events = read_csv(DATA / "project_events.csv")
    events_by_bill: dict[str, list[dict]] = defaultdict(list)
    for event in events:
        if event.get("boletin"):
            events_by_bill[event["boletin"]].append(event)

    existing = load_existing()
    current_bills = {p["boletin"] for p in projects if p.get("boletin")}
    corpus = {bill: row for bill, row in existing.items() if bill in current_bills}

    refreshed = 0
    reused = 0
    no_candidate = []
    errors = []

    for idx, project in enumerate(sorted(projects, key=lambda p: (p.get("fecha_ingreso", ""), p.get("boletin", ""))), start=1):
        bill = project.get("boletin", "")
        candidates = sorted(events_by_bill.get(bill, []), key=document_priority)
        candidate = next((e for e in candidates if document_priority(e) < 50), None)
        if not candidate:
            no_candidate.append(bill)
            continue
        source_url = candidate.get("documento_url", "")
        previous = corpus.get(bill)
        if previous and previous.get("source_url") == source_url and previous.get("cleaned_text"):
            reused += 1
            continue
        try:
            extracted = fetch_document(source_url)
            corpus[bill] = {
                "boletin": bill,
                "project_id": project.get("project_id", ""),
                "fecha_ingreso": project.get("fecha_ingreso", ""),
                "origen_iniciativa": project.get("origen_iniciativa", ""),
                "titulo": project.get("titulo", ""),
                "document_type": classify_document(candidate),
                "event_label": candidate.get("subetapa", ""),
                "source_url": source_url,
                **extracted,
                "synced_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
            }
            refreshed += 1
        except Exception as exc:  # noqa: BLE001
            errors.append({"boletin": bill, "error": str(exc)})
            if previous:
                corpus[bill] = previous
        if idx % 50 == 0:
            print(f"Textos {idx}/{len(projects)} · nuevos={refreshed} · reutilizados={reused} · errores={len(errors)}")

    OUT.mkdir(parents=True, exist_ok=True)
    ordered = [corpus[b] for b in sorted(corpus)]
    with TEXTS_PATH.open("w", encoding="utf-8") as handle:
        for row in ordered:
            handle.write(json.dumps(row, ensure_ascii=False) + "\n")

    index_fields = [
        "boletin", "project_id", "fecha_ingreso", "origen_iniciativa", "titulo",
        "document_type", "event_label", "source_url", "resolved_url", "mime_type",
        "extraction_method", "pages_total", "units_examined", "raw_chars", "cleaned_chars",
        "text_quality", "text_sha256", "synced_at",
    ]
    with INDEX_PATH.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=index_fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(ordered)

    by_origin = {}
    for origin in ("parlamentario", "ejecutivo"):
        expected = [p for p in projects if p.get("origen_iniciativa") == origin]
        rows = [r for r in ordered if r.get("origen_iniciativa") == origin]
        by_origin[origin] = {
            "projects": len(expected),
            "texts": len(rows),
            "rich_or_usable": sum(r.get("text_quality") in {"rica", "utilizable"} for r in rows),
        }

    report = {
        "projects": len(projects),
        "corpus_rows": len(ordered),
        "refreshed_this_run": refreshed,
        "reused_this_run": reused,
        "no_candidate_documents": no_candidate,
        "quality": dict(Counter(r.get("text_quality", "") for r in ordered)),
        "formats": dict(Counter(r.get("mime_type", "") for r in ordered)),
        "by_origin": by_origin,
        "errors": errors,
    }
    DIAG_PATH.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))

    coverage = len(ordered) / max(len(projects), 1)
    if coverage < 0.90:
        raise RuntimeError(f"Cobertura textual demasiado baja: {coverage:.1%}")
    if errors and len(errors) / max(len(projects), 1) > 0.05:
        raise RuntimeError(f"Demasiados errores de extracción: {len(errors)}/{len(projects)}")


if __name__ == "__main__":
    main()
