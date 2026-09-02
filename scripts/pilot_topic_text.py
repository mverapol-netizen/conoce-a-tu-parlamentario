from __future__ import annotations

import csv
import io
import json
import re
import unicodedata
from collections import Counter, defaultdict
from pathlib import Path

import requests
from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data" / "legislative" / "2026"
OUT = ROOT / "data" / "legislative" / "topic_pilot"
SAMPLE_PER_ORIGIN = 12
MAX_PDF_PAGES = 8
MAX_PREVIEW = 700

HTTP = requests.Session()
HTTP.headers.update({
    "User-Agent": "Mozilla/5.0 conoce-a-tu-parlamentario/topic-text-pilot",
    "Accept-Language": "es-CL,es;q=0.9",
})


def read_csv(path: Path) -> list[dict]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def norm(value: str) -> str:
    text = unicodedata.normalize("NFD", str(value or "").lower())
    text = "".join(ch for ch in text if unicodedata.category(ch) != "Mn")
    return re.sub(r"\s+", " ", text).strip()


def evenly_spaced(rows: list[dict], n: int) -> list[dict]:
    if len(rows) <= n:
        return rows[:]
    if n <= 1:
        return [rows[0]]
    idxs = sorted({round(i * (len(rows) - 1) / (n - 1)) for i in range(n)})
    return [rows[i] for i in idxs]


def document_priority(event: dict) -> int:
    label = norm(event.get("subetapa", ""))
    url = (event.get("documento_url") or "").strip()
    if not url:
        return 99
    if "texto del mensaje" in label:
        return 0
    if "texto del proyecto" in label:
        return 1
    if "ingreso de proyecto" in label:
        return 2
    return 50


def classify_document(event: dict) -> str:
    label = norm(event.get("subetapa", ""))
    if "texto del mensaje" in label:
        return "texto_mensaje"
    if "texto del proyecto" in label:
        return "texto_proyecto"
    if "ingreso de proyecto" in label:
        return "ingreso_proyecto"
    return "otro"


def extract_pdf(content: bytes) -> tuple[str, int, int]:
    reader = PdfReader(io.BytesIO(content))
    total = len(reader.pages)
    parts = []
    examined = min(total, MAX_PDF_PAGES)
    for page in reader.pages[:examined]:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            parts.append("")
    return "\n".join(parts), total, examined


def extract_docx(content: bytes) -> tuple[str, int]:
    document = Document(io.BytesIO(content))
    pieces = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            pieces.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    return "\n".join(pieces), len(document.paragraphs)


def extract_html(content: bytes, encoding: str | None) -> str:
    text = content.decode(encoding or "utf-8", errors="replace")
    soup = BeautifulSoup(text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    return re.sub(r"\s+", " ", soup.get_text(" ", strip=True)).strip()


def quality(chars: int) -> str:
    if chars >= 2000:
        return "rica"
    if chars >= 500:
        return "utilizable"
    if chars > 0:
        return "escasa"
    return "sin_texto"


def fetch_document(url: str) -> dict:
    response = HTTP.get(url, timeout=60, allow_redirects=True)
    status = response.status_code
    response.raise_for_status()
    content = response.content
    ctype = (response.headers.get("Content-Type") or "").lower()
    resolved = response.url
    method = ""
    pages_total = 0
    pages_examined = 0
    text = ""

    if content.lstrip().startswith(b"%PDF") or "application/pdf" in ctype:
        method = "pdf_text"
        text, pages_total, pages_examined = extract_pdf(content)
    elif "officedocument.wordprocessingml.document" in ctype:
        method = "docx_text"
        text, paragraphs = extract_docx(content)
        pages_total = 0
        pages_examined = paragraphs
    elif "html" in ctype or content.lstrip().startswith(b"<"):
        method = "html_text"
        text = extract_html(content, response.encoding)
    else:
        method = "unsupported"

    clean = re.sub(r"\s+", " ", text or "").strip()
    return {
        "http_status": status,
        "resolved_url": resolved,
        "mime_type": ctype.split(";", 1)[0],
        "extraction_method": method,
        "pages_total": pages_total,
        "pages_examined": pages_examined,
        "extracted_chars": len(clean),
        "text_quality": quality(len(clean)),
        "preview": clean[:MAX_PREVIEW],
    }


def main() -> None:
    projects = read_csv(DATA / "projects.csv")
    events = read_csv(DATA / "project_events.csv")
    events_by_bill: dict[str, list[dict]] = defaultdict(list)
    for event in events:
        if event.get("boletin"):
            events_by_bill[event["boletin"]].append(event)

    selected = []
    for origin in ("parlamentario", "ejecutivo"):
        group = sorted(
            [p for p in projects if p.get("origen_iniciativa") == origin],
            key=lambda p: (p.get("fecha_ingreso", ""), p.get("boletin", "")),
        )
        selected.extend(evenly_spaced(group, SAMPLE_PER_ORIGIN))

    rows = []
    for project in selected:
        candidates = sorted(events_by_bill.get(project["boletin"], []), key=document_priority)
        candidate = next((e for e in candidates if document_priority(e) < 50), None)
        row = {
            "boletin": project.get("boletin", ""),
            "project_id": project.get("project_id", ""),
            "fecha_ingreso": project.get("fecha_ingreso", ""),
            "origen_iniciativa": project.get("origen_iniciativa", ""),
            "titulo": project.get("titulo", ""),
            "document_type": "",
            "event_label": "",
            "source_url": "",
            "resolved_url": "",
            "http_status": "",
            "mime_type": "",
            "extraction_method": "",
            "pages_total": 0,
            "pages_examined": 0,
            "extracted_chars": 0,
            "text_quality": "sin_documento",
            "preview": "",
            "error": "",
        }
        if candidate:
            row["document_type"] = classify_document(candidate)
            row["event_label"] = candidate.get("subetapa", "")
            row["source_url"] = candidate.get("documento_url", "")
            try:
                row.update(fetch_document(row["source_url"]))
            except Exception as exc:  # noqa: BLE001
                row["error"] = str(exc)
                row["text_quality"] = "error"
        rows.append(row)
        print(f"{row['boletin']} · {row['origen_iniciativa']} · {row['text_quality']} · {row['extracted_chars']} chars")

    OUT.mkdir(parents=True, exist_ok=True)
    fields = list(rows[0].keys()) if rows else []
    with (OUT / "text_sources.csv").open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields)
        writer.writeheader()
        writer.writerows(rows)

    by_origin = {}
    for origin in ("parlamentario", "ejecutivo"):
        subset = [r for r in rows if r["origen_iniciativa"] == origin]
        by_origin[origin] = {
            "sample": len(subset),
            "with_candidate_document": sum(bool(r["source_url"]) for r in subset),
            "with_any_text": sum(int(r["extracted_chars"]) > 0 for r in subset),
            "usable_or_rich": sum(r["text_quality"] in {"utilizable", "rica"} for r in subset),
        }

    report = {
        "sample_projects": len(rows),
        "sample_per_origin_target": SAMPLE_PER_ORIGIN,
        "document_types": dict(Counter(r["document_type"] or "none" for r in rows)),
        "mime_types": dict(Counter(r["mime_type"] or "none" for r in rows)),
        "extraction_methods": dict(Counter(r["extraction_method"] or "none" for r in rows)),
        "quality": dict(Counter(r["text_quality"] for r in rows)),
        "by_origin": by_origin,
        "errors": [
            {"boletin": r["boletin"], "error": r["error"]}
            for r in rows if r["error"]
        ],
    }
    (OUT / "text_audit.json").write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    md = [
        "# Piloto de recuperación textual para clasificación temática",
        "",
        f"- Proyectos auditados: **{len(rows)}**",
        f"- Parlamentarios: **{by_origin['parlamentario']['sample']}**",
        f"- Ejecutivo: **{by_origin['ejecutivo']['sample']}**",
        f"- Con texto utilizable o rico: **{sum(r['text_quality'] in {'utilizable', 'rica'} for r in rows)}**",
        f"- Con algún texto: **{sum(int(r['extracted_chars']) > 0 for r in rows)}**",
        "",
        "La taxonomía temática no se aplica aún. Este piloto solo audita disponibilidad y calidad del texto fuente.",
    ]
    (OUT / "REPORT.md").write_text("\n".join(md) + "\n", encoding="utf-8")

    if not rows:
        raise RuntimeError("No se seleccionaron proyectos")
    if sum(bool(r["source_url"]) for r in rows) < int(len(rows) * 0.75):
        raise RuntimeError("Cobertura documental candidata inferior a 75%")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
